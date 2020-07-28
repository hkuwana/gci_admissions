# Python program to help GCI create letters to make it easier for them

from docx import Document
from docx.shared import Pt, RGBColor
import os
import pandas as pd

# Ensures that the file is read correctly. Each one corresponds to different parts of the excel sheet

df = pd.read_excel('Offer Letters and Excel/Mail Merge_anon.xlsx', sheet_name='Offer Letters 2020')
name_list = df['Name']
acceptance_status = [1, 'WL+', 'WL', 'P']


# ----------- Class ----------------

# ----------- Functions ------------

# Deletes the paragraph
def delete_paragraph(para):
    p = para._element
    p.getparent().remove(p)
    p._p = p._element = None

# Ensures the font are convereted into the write style (color, size, text, etc)
def font_converter(font_document, font_name, font_size, rgb_red, rgb_green, rgb_blue):
    font_style = font_document.styles['Normal'].font
    font_style.name = font_name
    font_style.size = Pt(font_size)
    font_style.color.rgb = RGBColor(rgb_red, rgb_green, rgb_blue)
    return font_style


# Make new directory to put all folders in
def check_directory():
    directory = "Student Offer Letters"
    accept = "/Accepted Students"
    waitliplus = "/Waitlist Plus Students"
    waitli = "/Waitlist Students"
    declin = "/Declined Students"
    if not os.path.exists(directory):
        os.mkdir(directory)
        os.mkdir(directory + accept)
        os.mkdir(directory + waitliplus)
        os.mkdir(directory +waitli)
        os.mkdir(directory + declin)
        print("Directory '% s' created" % directory)
    else:
        print("Directory '% s' already created" % directory)


# Function to separate the students type in a list
def seperate_students(student_list, student_acceptance_status):
    accepted = []
    waitlist_plus = []
    waitlist = []
    declined = []
    for student in range(len(student_list)):
        if isinstance((df['F1'][student]), int):
            accepted.append(student)
        elif (df['F1'][student]) == student_acceptance_status[1]:
            waitlist_plus.append(student)
        elif (df['F1'][student]) == student_acceptance_status[2]:
            waitlist.append(student)
        elif (df['F1'][student]) == student_acceptance_status[3]:
            declined.append(student)
        else:
            pass
    return accepted, waitlist_plus, waitlist, declined


# Creates and writes the Accepted Students into the proper file. Accepted takes in the array position
# of the accepted students, and excel_list takes in the array itself

def accepted_letters(accepted, excel_list):
    for i in accepted:
        personal_info = excel_list.loc[i, :]
        document = Document('Offer Letters and Excel/Fellowship Offer Letter Template.docx')
        font = font_converter(document, 'Arial', 10, 0, 0, 0)
        for paragraph in document.paragraphs:
            if 'SURNAME' in paragraph.text:
                paragraph.style = document.styles['Normal']
                paragraph.text = 'Dear Ms/Mr. %s,' % excel_list.loc[i, 'Name']
            elif 'scholarship to cover the Fellowship tuition and a (partial) travel grant (see details below)' \
                    in paragraph.text and excel_list.loc[i, 'TUITION'].casefold() is not 'full paying':
                if 'scholarship' or 'funder' in excel_list.loc[i, 'TUITION'].casefold():
                    paragraph.text = 'Congratulations! On behalf of Global Citizens Initiative (GCI), we are pleased ' \
                                     'to offer you admission to the GCI Fellowship 2020. You have also been awarded a' \
                                     ' full $5,900 scholarship to cover the Fellowship tuition and a travel grant ' \
                                     '(see details below)!'
                elif not excel_list.loc[i, 'TUITION'].casefold():
                    paragraph.text = 'Congratulations! On behalf of Global Citizens Initiative (GCI), we are pleased ' \
                                     'to offer you admission to the GCI Fellowship 2020. You have also been awarded ' \
                                     'a partial scholarship to cover the Fellowship tuition and a partial travel ' \
                                     'grant (see details below)!'
                else:
                    paragraph.text = 'Congratulations! On behalf of Global Citizens Initiative (GCI), we are pleased ' \
                                     'to offer you admission to the GCI Fellowship 2020.'
            elif 'As part of your Commitment to Enroll, you must complete the Visa Appointment Information' \
                    in paragraph.text and 'no' in excel_list.loc[i, 'Visa Required?'].casefold():
                delete_paragraph(paragraph)
            elif 'Congratulations! You have been awarded a travel grant!' in paragraph.text and 'full paying' or 'reimburse' in excel_list.loc[i, 'TRAVEL'].casefold():
                delete_paragraph(paragraph)
            elif 'Congratulations! You have been awarded a partial travel grant!' in paragraph.text and 'full paying' or 'funder'  in excel_list.loc[i, 'TRAVEL'].casefold():
                delete_paragraph(paragraph)
            elif 'Make the first payment of ' in paragraph.text and 'full paying' not in excel_list.loc[i, 'TUITION' ].casefold():
                delete_paragraph(paragraph)
            elif 'Make the second payment of ' in paragraph.text and 'full paying' not in excel_list.loc[i, 'TUITION'].casefold():
                delete_paragraph(paragraph)
            elif 'Payment options: via wire transfer' in paragraph.text and 'full paying' not in excel_list.loc[i, 'TUITION'].casefold():
                delete_paragraph(paragraph)
            elif 'Enrollment Cancellation Policy' in paragraph.text and 'full paying' not in excel_list.loc[i, 'TUITION'].casefold():
                delete_paragraph(paragraph)
        document.save('Student Offer Letters/Accepted Students/%s.docx' % excel_list.loc[i, 'Name'])

def waitlist_plus_letters(waitlist_plus, excel_list):
    for i in waitlist_plus:
        personal_info = excel_list.loc[i, :]
        document = Document('Offer Letters and Excel/Fellowship Waitlist+ Letter Template.docx')
        font = font_converter(document, 'Arial', 10, 0, 0, 0)
        for paragraph in document.paragraphs:
            if 'SURNAME' in paragraph.text:
                paragraph.style = document.styles['Normal']
                paragraph.text = 'Dear Ms/Mr. %s,' % excel_list.loc[i, 'Name']
        document.save('Student Offer Letters/Waitlist Plus Students/%s.docx' % excel_list.loc[i, 'Name'])


def waitlist_letters(waitlist, excel_list):
    for i in waitlist:
        personal_info = excel_list.loc[i, :]
        document = Document('Offer Letters and Excel/Fellowship Waitlist Letter Template.docx')
        font = font_converter(document, 'Arial', 10, 0, 0, 0)
        for paragraph in document.paragraphs:
            if 'SURNAME' in paragraph.text:
                paragraph.style = document.styles['Normal']
                paragraph.text = 'Dear Ms/Mr. %s,' % excel_list.loc[i, 'Name']
        document.save('Student Offer Letters/Waitlist Students/%s.docx' % excel_list.loc[i, 'Name'])


def declined_letters(declined, excel_list):
    for i in declined:
        personal_info = excel_list.loc[i, :]
        document = Document('Offer Letters and Excel/Fellowship Declined Letter Template.docx')
        font = font_converter(document, 'Arial', 10, 0, 0, 0)
        for paragraph in document.paragraphs:
            if 'Applicant' in paragraph.text:
                paragraph.style = document.styles['Normal']
                paragraph.text = 'Dear Ms/Mr. %s,' % excel_list.loc[i, 'Name']
        document.save('Student Offer Letters/Declined Students/%s.docx' % excel_list.loc[i, 'Name'])

def create_letters():
    accepted_students, waitlist_plus_students, waitlist_students, declined_students = \
        seperate_students(name_list, acceptance_status)
    check_directory()
    accepted_letters(accepted_students, df)
    waitlist_plus_letters(waitlist_plus_students, df)
    waitlist_letters(waitlist_students, df)
    declined_letters(declined_students, df)

# Checks every student to make sure the students are in the right acceptance status.
# Need to make sure that the accepted students are going to be A rather than integers, but if kept
# as integers, that's okay, too

create_letters()

